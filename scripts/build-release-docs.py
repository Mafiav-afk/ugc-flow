from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCES = [
    "UGC-Flow-软件操作说明",
    "UGC-Flow-API接入说明",
    "UGC-Flow-视频讲解脚本",
]

BLUE = RGBColor(0xD9, 0x42, 0x36)
DARK = RGBColor(0x20, 0x29, 0x2F)
MUTED = RGBColor(0x68, 0x72, 0x7A)
LIGHT = "F5F6F8"
FONT = "Hiragino Sans GB"

def set_font(run, size=11, bold=False, color=DARK, italic=False, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def shade_paragraph(paragraph, fill=LIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, 9, color=MUTED)

def configure(doc, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [("Heading 1", 15, 13, 6), ("Heading 2", 12.5, 10, 5), ("Heading 3", 11, 7, 3)]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if name != "Heading 3" else DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("UGC FLOW  ·  发布资料")
    set_font(run, 8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("UGC FLOW OPERATOR GUIDE")
    set_font(r, 9, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title)
    set_font(r, 23, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Mac 版 1.3.0  ·  Codex / Claude 智能体 + 图片与视频 API")
    set_font(r, 10.5, color=MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    ppr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D94236")
    borders.append(bottom)
    ppr.append(borders)

def add_inline_runs(paragraph, text, base_size=10):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 9.2, color=RGBColor(0x8A, 0x2D, 0x25), name=FONT)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, base_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, base_size)

def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "0")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id

def add_numbered(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    add_inline_runs(p, text)
    return p

def build(source_name):
    md_path = DOCS / f"{source_name}.md"
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    doc = Document()
    configure(doc, title)
    in_code = False
    code_lines = []
    in_numbered_list = False
    num_id = None
    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                shade_paragraph(p)
                r = p.add_run("\n".join(code_lines))
                set_font(r, 8.3, color=DARK, name=FONT)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        numbered = bool(re.match(r"^\d+\. ", line))
        if not numbered:
            in_numbered_list = False
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 1")
        elif numbered:
            if not in_numbered_list:
                num_id = new_numbering_id(doc)
                in_numbered_list = True
            p = add_numbered(doc, re.sub(r"^\d+\. ", "", line), num_id)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, "FFF3F1")
            add_inline_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line.replace("  ", " "))
    out = DOCS / f"{source_name}.docx"
    doc.save(out)
    return out

if __name__ == "__main__":
    DOCS.mkdir(exist_ok=True)
    for source in SOURCES:
        print(build(source))
